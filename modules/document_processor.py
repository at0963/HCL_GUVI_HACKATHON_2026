"""
Document Processor Module
Handles PDF, DOCX, DOC, and TXT file extraction
"""
import os
import io
from pathlib import Path
from typing import Dict, Optional, Tuple
import PyPDF2
import pdfplumber
from docx import Document
import re


class DocumentProcessor:
    """Process various document formats and extract text"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    
    def process_document(self, file_path: str) -> Dict:
        """
        Main method to process any supported document type
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dict containing extracted text, metadata, and status
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {
                "success": False,
                "error": "File not found",
                "text": "",
                "metadata": {}
            }
        
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_formats:
            return {
                "success": False,
                "error": f"Unsupported format: {extension}",
                "text": "",
                "metadata": {}
            }
        
        try:
            if extension == '.pdf':
                return self._process_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                return self._process_docx(file_path)
            elif extension == '.txt':
                return self._process_txt(file_path)
        except Exception as e:
            return {
                "success": False,
                "error": f"Error processing document: {str(e)}",
                "text": "",
                "metadata": {}
            }
    
    def _process_pdf(self, file_path: Path) -> Dict:
        """Extract text from PDF using multiple methods"""
        text = ""
        metadata = {
            "pages": 0,
            "extraction_method": "pdfplumber"
        }
        
        try:
            # Try pdfplumber first (better for structured text)
            with pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            
            # If pdfplumber fails, fallback to PyPDF2
            if not text.strip():
                metadata["extraction_method"] = "PyPDF2"
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata["pages"] = len(pdf_reader.pages)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
            
            # Clean extracted text
            text = self._clean_text(text)
            
            return {
                "success": True,
                "text": text,
                "metadata": metadata,
                "file_name": file_path.name,
                "file_type": "PDF"
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": f"PDF processing error: {str(e)}",
                "text": "",
                "metadata": metadata
            }
    
    def _process_docx(self, file_path: Path) -> Dict:
        """Extract text from DOCX/DOC files"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"
            
            # Clean extracted text
            text = self._clean_text(text)
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            }
            
            return {
                "success": True,
                "text": text,
                "metadata": metadata,
                "file_name": file_path.name,
                "file_type": "DOCX"
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": f"DOCX processing error: {str(e)}",
                "text": "",
                "metadata": {}
            }
    
    def _process_txt(self, file_path: Path) -> Dict:
        """Extract text from TXT files"""
        try:
            # Try multiple encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not text:
                raise ValueError("Could not decode text file with any supported encoding")
            
            # Clean extracted text
            text = self._clean_text(text)
            
            metadata = {
                "lines": len(text.split('\n')),
                "encoding": encoding
            }
            
            return {
                "success": True,
                "text": text,
                "metadata": metadata,
                "file_name": file_path.name,
                "file_type": "TXT"
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": f"TXT processing error: {str(e)}",
                "text": "",
                "metadata": {}
            }
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that cause issues
        text = text.replace('\x00', '')
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """
        Extract common contract sections
        
        Returns:
            Dict with section names as keys and content as values
        """
        sections = {}
        
        # Common section headers in contracts
        section_patterns = [
            r'(?:^|\n)(?:ARTICLE|SECTION|CLAUSE)\s*\d+[\.:]\s*(.+?)(?=(?:\n(?:ARTICLE|SECTION|CLAUSE)\s*\d+)|$)',
            r'(?:^|\n)(\d+\.\s*.+?)(?=(?:\n\d+\.)|$)',
            r'(?:^|\n)([A-Z][A-Z\s]{10,})(?=\n)',  # All caps headers
        ]
        
        for pattern in section_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
            for i, match in enumerate(matches):
                section_name = match.group(1).strip() if match.groups() else f"Section {i+1}"
                sections[section_name] = match.group(0).strip()
        
        return sections if sections else {"Full Document": text}
    
    def process_uploaded_file(self, uploaded_file) -> Dict:
        """
        Process file from Streamlit file uploader
        
        Args:
            uploaded_file: Streamlit UploadedFile object
            
        Returns:
            Dict containing extracted text and metadata
        """
        import tempfile
        
        try:
            # Store original filename
            original_filename = uploaded_file.name
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(uploaded_file.name).suffix) as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_path = tmp_file.name
            
            # Process the temporary file
            result = self.process_document(tmp_path)
            
            # Replace temp filename with original filename
            if result.get("success") and "file_name" in result:
                result["file_name"] = original_filename
            
            # Clean up
            os.unlink(tmp_path)
            
            return result
        
        except Exception as e:
            return {
                "success": False,
                "error": f"Upload processing error: {str(e)}",
                "text": "",
                "metadata": {}
            }
